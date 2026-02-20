"""
Word Document Loader.

Extracts text from .docx files including:
- Paragraph text with heading detection
- Table content in a structured format
- Metadata (title, author if available)

Why custom instead of LangChain's Docx2txtLoader?
- We preserve heading structure (important for chunking quality)
- We extract tables properly (critical for audit documents)
- We add richer metadata
"""
from langchain_core.documents import Document
from docx import Document as DocxDocument
import logging

logger = logging.getLogger(__name__)


def load_docx(file_path: str) -> list[Document]:
    """
    Load a Word document and return LangChain Document objects.

    The document is returned as a single Document (not split by page,
    because Word documents don't have natural page breaks in the file).
    Chunking happens later in the pipeline.

    Args:
        file_path: Path to the .docx file

    Returns:
        List containing one Document with the full text
    """
    try:
        doc = DocxDocument(file_path)
        sections = []

        # Extract paragraphs with heading awareness
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Detect headings and mark them (helps chunking later)
            style_name = paragraph.style.name if paragraph.style else ""
            if "Heading" in style_name:
                # Add a marker so the text splitter can use headings as boundaries
                sections.append(f"\n## {text}\n")
            else:
                sections.append(text)

        # Extract tables
        for table in doc.tables:
            table_text = "\n[TABLE DATA]\n"
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                for row in table.rows[1:]:
                    row_text = " | ".join(
                        f"{h}: {cell.text.strip()}"
                        for h, cell in zip(headers, row.cells)
                    )
                    table_text += row_text + "\n"
            sections.append(table_text)

        full_text = "\n".join(sections)

        if full_text.strip():
            documents = [Document(
                page_content=full_text.strip(),
                metadata={
                    "source": file_path,
                    "file_type": "docx",
                    "has_tables": len(doc.tables) > 0,
                    "paragraph_count": len(doc.paragraphs),
                }
            )]
            logger.info(f"Loaded DOCX: {file_path} — {len(doc.paragraphs)} paragraphs")
            return documents

        return []

    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        raise
